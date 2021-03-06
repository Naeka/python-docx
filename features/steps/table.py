# encoding: utf-8

"""
Step implementations for table-related features
"""

from __future__ import (
    absolute_import, division, print_function, unicode_literals
)

from behave import given, then, when

from docx import Document
from docx.shared import Inches
from docx.table import _Column, _Columns, _Row, _Rows

from helpers import test_docx


# given ===================================================

@given('a 2 x 2 table')
def given_a_2x2_table(context):
    context.table_ = Document().add_table(rows=2, cols=2)


@given('a 3x3 table having {span_state}')
def given_a_3x3_table_having_span_state(context, span_state):
    table_idx = {
        'only uniform cells': 0,
        'a horizontal span':  1,
        'a vertical span':    2,
        'a combined span':    3,
    }[span_state]
    document = Document(test_docx('tbl-cell-access'))
    context.table_ = document.tables[table_idx]


@given('a column collection having two columns')
def given_a_column_collection_having_two_columns(context):
    docx_path = test_docx('blk-containing-table')
    document = Document(docx_path)
    context.columns = document.tables[0].columns


@given('a row collection having two rows')
def given_a_row_collection_having_two_rows(context):
    docx_path = test_docx('blk-containing-table')
    document = Document(docx_path)
    context.rows = document.tables[0].rows


@given('a table')
def given_a_table(context):
    context.table_ = Document().add_table(rows=2, cols=2)


@given('a table cell having a width of {width}')
def given_a_table_cell_having_a_width_of_width(context, width):
    table_idx = {'no explicit setting': 0, '1 inch': 1, '2 inches': 2}[width]
    document = Document(test_docx('tbl-props'))
    table = document.tables[table_idx]
    cell = table.cell(0, 0)
    context.cell = cell


@given('a table column having a width of {width_desc}')
def given_a_table_having_a_width_of_width_desc(context, width_desc):
    col_idx = {
        'no explicit setting': 0,
        '1440':                1,
    }[width_desc]
    docx_path = test_docx('tbl-col-props')
    document = Document(docx_path)
    context.column = document.tables[0].columns[col_idx]


@given('a table having an applied style')
def given_a_table_having_an_applied_style(context):
    docx_path = test_docx('tbl-having-applied-style')
    document = Document(docx_path)
    context.table_ = document.tables[0]


@given('a table having an autofit layout of {autofit}')
def given_a_table_having_an_autofit_layout_of_autofit(context, autofit):
    tbl_idx = {
        'no explicit setting': 0,
        'autofit':             1,
        'fixed':               2,
    }[autofit]
    document = Document(test_docx('tbl-props'))
    context.table_ = document.tables[tbl_idx]


@given('a table having two columns')
def given_a_table_having_two_columns(context):
    docx_path = test_docx('blk-containing-table')
    document = Document(docx_path)
    # context.table is used internally by behave, underscore added
    # to distinguish this one
    context.table_ = document.tables[0]


@given('a table having two rows')
def given_a_table_having_two_rows(context):
    docx_path = test_docx('blk-containing-table')
    document = Document(docx_path)
    context.table_ = document.tables[0]


# when =====================================================

@when('I add a column to the table')
def when_add_column_to_table(context):
    table = context.table_
    context.column = table.add_column()


@when('I add a row to the table')
def when_add_row_to_table(context):
    table = context.table_
    context.row = table.add_row()


@when('I apply a style to the table')
def when_apply_style_to_table(context):
    table = context.table_
    table.style = 'LightShading-Accent1'


@when('I merge from cell {origin} to cell {other}')
def when_I_merge_from_cell_origin_to_cell_other(context, origin, other):
    def cell(table, idx):
        row, col = idx // 3, idx % 3
        return table.cell(row, col)
    a_idx, b_idx = int(origin) - 1, int(other) - 1
    table = context.table_
    a, b = cell(table, a_idx), cell(table, b_idx)
    a.merge(b)


@when('I set the cell width to {width}')
def when_I_set_the_cell_width_to_width(context, width):
    new_value = {'1 inch': Inches(1)}[width]
    context.cell.width = new_value


@when('I set the column width to {width_emu}')
def when_I_set_the_column_width_to_width_emu(context, width_emu):
    new_value = None if width_emu == 'None' else int(width_emu)
    context.column.width = new_value


@when('I set the table autofit to {setting}')
def when_I_set_the_table_autofit_to_setting(context, setting):
    new_value = {'autofit': True, 'fixed': False}[setting]
    table = context.table_
    table.autofit = new_value


# then =====================================================

@then('I can access a collection column by index')
def then_can_access_collection_column_by_index(context):
    columns = context.columns
    for idx in range(2):
        column = columns[idx]
        assert isinstance(column, _Column)


@then('I can access a collection row by index')
def then_can_access_collection_row_by_index(context):
    rows = context.rows
    for idx in range(2):
        row = rows[idx]
        assert isinstance(row, _Row)


@then('I can access the column collection of the table')
def then_can_access_column_collection_of_table(context):
    table = context.table_
    columns = table.columns
    assert isinstance(columns, _Columns)


@then('I can access the row collection of the table')
def then_can_access_row_collection_of_table(context):
    table = context.table_
    rows = table.rows
    assert isinstance(rows, _Rows)


@then('I can get the table style name')
def then_can_get_table_style_name(context):
    table = context.table_
    msg = "got '%s'" % table.style
    assert table.style == 'LightShading-Accent1', msg


@then('I can iterate over the column collection')
def then_can_iterate_over_column_collection(context):
    columns = context.columns
    actual_count = 0
    for column in columns:
        actual_count += 1
        assert isinstance(column, _Column)
    assert actual_count == 2


@then('I can iterate over the row collection')
def then_can_iterate_over_row_collection(context):
    rows = context.rows
    actual_count = 0
    for row in rows:
        actual_count += 1
        assert isinstance(row, _Row)
    assert actual_count == 2


@then('table.cell({row}, {col}).text is {expected_text}')
def then_table_cell_row_col_text_is_text(context, row, col, expected_text):
    table = context.table_
    row_idx, col_idx = int(row), int(col)
    cell_text = table.cell(row_idx, col_idx).text
    assert cell_text == expected_text, 'got %s' % cell_text


@then('the column cells text is {expected_text}')
def then_the_column_cells_text_is_expected_text(context, expected_text):
    table = context.table_
    cells_text = ' '.join(c.text for col in table.columns for c in col.cells)
    assert cells_text == expected_text, 'got %s' % cells_text


@then('the length of the column collection is 2')
def then_len_of_column_collection_is_2(context):
    columns = context.table_.columns
    assert len(columns) == 2


@then('the length of the row collection is 2')
def then_len_of_row_collection_is_2(context):
    rows = context.table_.rows
    assert len(rows) == 2


@then('the new column has 2 cells')
def then_new_column_has_2_cells(context):
    assert len(context.column.cells) == 2


@then('the new row has 2 cells')
def then_new_row_has_2_cells(context):
    assert len(context.row.cells) == 2


@then('the reported autofit setting is {autofit}')
def then_the_reported_autofit_setting_is_autofit(context, autofit):
    expected_value = {'autofit': True, 'fixed': False}[autofit]
    table = context.table_
    assert table.autofit is expected_value


@then('the reported column width is {width_emu}')
def then_the_reported_column_width_is_width_emu(context, width_emu):
    expected_value = None if width_emu == 'None' else int(width_emu)
    assert context.column.width == expected_value, (
        'got %s' % context.column.width
    )


@then('the reported width of the cell is {width}')
def then_the_reported_width_of_the_cell_is_width(context, width):
    expected_width = {'None': None, '1 inch': Inches(1)}[width]
    actual_width = context.cell.width
    assert actual_width == expected_width, (
        'expected %s, got %s' % (expected_width, actual_width)
    )


@then('the row cells text is {encoded_text}')
def then_the_row_cells_text_is_expected_text(context, encoded_text):
    expected_text = encoded_text.replace('\\', '\n')
    table = context.table_
    cells_text = ' '.join(c.text for row in table.rows for c in row.cells)
    assert cells_text == expected_text, 'got %s' % cells_text


@then('the table style matches the name I applied')
def then_table_style_matches_name_applied(context):
    table = context.table_
    tmpl = "table.style doesn't match, got '%s'"
    assert table.style == 'LightShading-Accent1', tmpl % table.style


@then('the table has {count} columns')
def then_table_has_count_columns(context, count):
    column_count = int(count)
    columns = context.table_.columns
    assert len(columns) == column_count


@then('the table has {count} rows')
def then_table_has_count_rows(context, count):
    row_count = int(count)
    rows = context.table_.rows
    assert len(rows) == row_count


@then('the width of cell {n_str} is {inches_str} inches')
def then_the_width_of_cell_n_is_x_inches(context, n_str, inches_str):
    def _cell(table, idx):
        row, col = idx // 3, idx % 3
        return table.cell(row, col)
    idx, inches = int(n_str) - 1, float(inches_str)
    cell = _cell(context.table_, idx)
    assert cell.width == Inches(inches), 'got %s' % cell.width.inches
